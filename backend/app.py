import os
import json
import re
import io
import google.generativeai as genai
from fastapi import FastAPI, Request # CHANGED: Imported FastAPI and Request
from fastapi.responses import RedirectResponse, JSONResponse # CHANGED: Imported FastAPI responses
from fastapi.middleware.cors import CORSMiddleware # CHANGED: Imported FastAPI CORS
from starlette.middleware.sessions import SessionMiddleware # CHANGED: Imported SessionMiddleware
from dotenv import load_dotenv
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
# RENAMED: Renamed 'Request' to 'GoogleAuthRequest' to avoid conflict with FastAPI's 'Request'
from google.auth.transport.requests import Request as GoogleAuthRequest 
from google_auth_oauthlib.flow import Flow 

import uvicorn # ADDED: For running the FastAPI server
import datetime # ADDED: This import was used in the grading logic

from google.cloud import vision # ADDED: Google Vision API client

# Document processing
from docx import Document # ADDED: For Word document text extraction

# --- LOCAL CPU MINILM MODEL IMPORTS ---
import torch
import numpy as np
from sentence_transformers import SentenceTransformer, util 

# --- 1. INITIAL CONFIGURATION ---
# Load .env from the same directory as this script
env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
load_dotenv(dotenv_path=env_path)
print(f"📁 Loading environment from: {env_path}")

# --- 2. ADD VISION API AUTHENTICATION ---
# This tells the script to use your JSON key file.
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "service-account-key.json"

app = FastAPI() # CHANGED: Initialized FastAPI


# ADDED: SessionMiddleware for session support, equivalent to Flask's app.secret_key
# Note: Using os.urandom() means sessions will be invalidated on every server restart.
# For production, use a static key, e.g., os.getenv("SESSION_SECRET_KEY").
app.add_middleware(SessionMiddleware, secret_key=os.urandom(24))

# CHANGED: Replaced Flask-CORS with FastAPI's CORSMiddleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"], # Explicitly adding methods
    allow_headers=["*"], # Explicitly adding headers
)

os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
CLIENT_SECRETS_FILE = 'client_secret.json'

# ALL SCOPES KEPT AS PER YOUR INSTRUCTION
# *** IMPORTANT: ADDED NEW SCOPES HERE ***
SCOPES = [ 
    'https://www.googleapis.com/auth/classroom.courses.readonly', 
    'https://www.googleapis.com/auth/classroom.coursework.students', 
    'https://www.googleapis.com/auth/classroom.course-work.readonly', 
    'https://www.googleapis.com/auth/classroom.profile.emails', # NEW SCOPE
    'https://www.googleapis.com/auth/classroom.profile.photos', # NEW SCOPE (Optional, but often included with emails)
    'https://www.googleapis.com/auth/drive.readonly',
    'https://www.googleapis.com/auth/classroom.coursework.me',
    'https://www.googleapis.com/auth/classroom.rosters.readonly',
    'https://www.googleapis.com/auth/classroom.student-submissions.students.readonly',
    'https://www.googleapis.com/auth/drive.file',  # NEW: Create files in Drive
    'https://www.googleapis.com/auth/spreadsheets',  # NEW: Create and edit Google Sheets
    'https://www.googleapis.com/auth/documents'  # For Google Docs API (answer key saving)
]

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# --- MongoDB Setup ---
from pymongo import MongoClient
from bson import ObjectId
import certifi

MONGO_URI = os.getenv("MONGO_URI")
print(f"🔍 MONGO_URI loaded: {'✅ Found' if MONGO_URI else '❌ Missing'}")
if MONGO_URI:
    print(f"🔗 MongoDB URI: {MONGO_URI[:20]}...{MONGO_URI[-20:]}")  # Show first/last 20 chars for security

# Initialize MongoDB connection
try:
    mongo_client = MongoClient(
        MONGO_URI, 
        tls=True,
        tlsAllowInvalidCertificates=True,  # Bypass SSL verification
        serverSelectionTimeoutMS=10000,  # Increased timeout
        connectTimeoutMS=10000,
        socketTimeoutMS=10000
    )
    db = mongo_client['gradepilot']
    
    # Collections
    grades_collection = db['grades']
    students_collection = db['students']
    
    # Create indexes for faster queries
    grades_collection.create_index([("course_id", 1)])
    grades_collection.create_index([("assignment_id", 1)])
    grades_collection.create_index([("student_name", 1)])
    grades_collection.create_index([("timestamp", -1)])
    grades_collection.create_index([("course_id", 1), ("assignment_id", 1)])
    grades_collection.create_index([("student_name", 1), ("course_id", 1)])
    
    students_collection.create_index([("student_name", 1)])
    students_collection.create_index([("course_id", 1)])
    
    print("✅ MongoDB connected successfully!")
    print(f"📊 Database: {db.name}")
    
except Exception as e:
    print(f"⚠️ MongoDB connection failed: {e}")
    print("⚠️ Will fall back to in-memory storage")
    mongo_client = None
    db = None
    grades_collection = None
    students_collection = None

# --- Global storage for graded assignments (in-memory fallback) ---
graded_assignments_history = []


# --- LOCAL CPU MINILM MODEL SETUP ---
# Initialize MiniLM model for semantic similarity grading (runs on local CPU/GPU)
MINILM_MODEL = None
MINILM_DEVICE = None

def initialize_minilm_model():
    """
    Initialize the MiniLM model on the best available device.
    First tries to load fine-tuned model from ./minilm-finetuned-grading
    Falls back to base model if fine-tuned version not found.
    Will try GPU first (if available), otherwise fall back to CPU.
    """
    global MINILM_MODEL, MINILM_DEVICE
    
    if MINILM_MODEL is not None:
        return MINILM_MODEL
    
    # Check for fine-tuned model first
    import os
    FINETUNED_PATH = "./minilm-finetuned-grading"
    BASE_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
    
    if os.path.exists(FINETUNED_PATH):
        MODEL_NAME = FINETUNED_PATH
        print("📚 Found fine-tuned model for answer grading")
    else:
        MODEL_NAME = BASE_MODEL_NAME
        print("📚 Using base MiniLM model (run 'python finetune_minilm.py' to create fine-tuned version)")
    
    # Auto-select device (CPU or best available GPU)
    if torch.cuda.is_available():
        device_count = torch.cuda.device_count()
        print(f"✅ CUDA available with {device_count} GPU(s)")
        # Try to use the first available GPU
        MINILM_DEVICE = "cuda:0"
        print(f"🎯 Using GPU: {MINILM_DEVICE}")
    else:
        MINILM_DEVICE = "cpu"
        print("ℹ️ CUDA not available, using CPU")
    
    print(f"📥 Loading model on {MINILM_DEVICE}...")
    
    try:
        MINILM_MODEL = SentenceTransformer(MODEL_NAME, device=MINILM_DEVICE)
        model_type = "FINE-TUNED" if MODEL_NAME == FINETUNED_PATH else "BASE"
        print(f"✅ {model_type} MiniLM model loaded successfully on {MINILM_DEVICE}")
    except Exception as e:
        print(f"⚠️ Failed to load on {MINILM_DEVICE}: {e}")
        if MINILM_DEVICE != "cpu":
            print("🔄 Retrying on CPU...")
            MINILM_DEVICE = "cpu"
            try:
                MINILM_MODEL = SentenceTransformer(MODEL_NAME, device=MINILM_DEVICE)
                model_type = "FINE-TUNED" if MODEL_NAME == FINETUNED_PATH else "BASE"
                print(f"✅ {model_type} MiniLM model loaded successfully on CPU")
            except Exception as cpu_error:
                print(f"❌ Failed to load MiniLM model on CPU: {cpu_error}")
                print("⚠️ Will use Gemini-only grading")
                MINILM_MODEL = None
    
    return MINILM_MODEL


def get_minilm_semantic_score(teacher_answer, student_answer):
    """
    Calculate semantic similarity between teacher and student answers using MiniLM model.
    
    Args:
        teacher_answer (str): Reference/correct answer
        student_answer (str): Student's submitted answer
        
    Returns:
        float or None: Cosine similarity score (0.0-1.0), or None if model unavailable
    """
    try:
        model = initialize_minilm_model()
        
        if model is None:
            print("⚠️ MiniLM model not available")
            return None
        
        # Encode both answers
        teacher_emb = model.encode(teacher_answer, convert_to_numpy=True)
        student_emb = model.encode(student_answer, convert_to_numpy=True)
        
        # Calculate cosine similarity
        similarity = float(np.dot(teacher_emb, student_emb) / (np.linalg.norm(teacher_emb) * np.linalg.norm(student_emb)))
        
        return similarity
    
    except Exception as e:
        print(f"⚠️ Error calculating MiniLM similarity: {e}")
        return None


def normalize_minilm_score_to_grade(similarity_score):
    """
    Convert MiniLM similarity score (0.3-0.85) to a grade (0-100) using min-max normalization.
    
    Formula: Mapped Score = ((x - 0.3) / (0.85 - 0.3)) × 100
    
    Where:
    - x = similarity score
    - 0.3 = minimum of scale
    - 0.85 = maximum of scale
    
    This maps:
    - 0.30 similarity → 0/100 grade
    - 0.575 similarity → 50/100 grade  
    - 0.85 similarity → 100/100 grade
    
    Args:
        similarity_score (float): Cosine similarity from MiniLM model (0.0-1.0)
        
    Returns:
        int: Grade out of 100
    """
    if similarity_score is None:
        return None
    
    # Clamp to valid range [0.3, 0.85]
    score = max(0.3, min(0.85, similarity_score))
    
    # Min-max normalization: map [0.3, 0.85] to [0, 100]
    mapped_score = ((score - 0.3) / (0.85 - 0.3)) * 100
    
    return int(round(mapped_score))


def compare_minilm_and_gemini_grades(minilm_grade, gemini_grade, threshold=15):
    """
    Compare MiniLM and Gemini grades to determine confidence.
    
    Args:
        minilm_grade (int): Grade from MiniLM model (0-100)
        gemini_grade (int): Grade from Gemini AI (0-100)
        threshold (int): Acceptable difference (default: 15 points)
        
    Returns:
        dict: Analysis with confidence level and recommendation
    """
    if minilm_grade is None:
        return {
            'confidence': 'medium',
            'final_grade': gemini_grade,
            'method': 'gemini_only',
            'difference': None,
            'recommendation': 'MiniLM model unavailable, using Gemini only'
        }
    
    difference = abs(minilm_grade - gemini_grade)
    
    if difference <= threshold:
        # Grades agree - high confidence
        # Use average of both for more balanced result
        final_grade = round((minilm_grade + gemini_grade) / 2)
        return {
            'confidence': 'high',
            'final_grade': final_grade,
            'method': 'hybrid_average',
            'difference': difference,
            'minilm_grade': minilm_grade,
            'gemini_grade': gemini_grade,
            'recommendation': f'Both models agree (diff: {difference}pts). Using average.'
        }
    else:
        # Grades disagree - lower confidence, prefer Gemini for detailed reasoning
        return {
            'confidence': 'medium',
            'final_grade': gemini_grade,
            'method': 'gemini_preferred',
            'difference': difference,
            'minilm_grade': minilm_grade,
            'gemini_grade': gemini_grade,
            'recommendation': f'Models disagree by {difference}pts. Using Gemini (better reasoning).'
        }
# --- END LOCAL CPU MINILM MODEL SETUP ---


# --- 2. HELPER FUNCTIONS (No changes here, keeping for full context) ---

def credentials_to_dict(credentials):
    """Converts a Google Credentials object to a dictionary for session storage."""
    # NO CHANGE: This function is framework-independent.
    return {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }

# MODIFIED: Function now requires the 'request' object to access the session
def get_google_service(service_name, version, request: Request): 
    """
    Builds and returns an authorized Google API service object (e.g., Classroom, Drive).
    Handles refreshing expired access tokens using the refresh token.
    
    MODIFIED: Now takes 'request: Request' as an argument to access session data.
    """
    # CHANGED: 'session' is now 'request.session'
    if 'credentials' not in request.session:
        return None # User not authenticated

    # CHANGED: 'session' is now 'request.session'
    creds_data = request.session['credentials']
    creds = Credentials(
        token=creds_data['token'],
        refresh_token=creds_data['refresh_token'],
        token_uri=creds_data['token_uri'],
        client_id=creds_data['client_id'],
        client_secret=creds_data['client_secret'],
        scopes=creds_data['scopes']
    )

    if not creds.valid:
        if creds.expired and creds.refresh_token:
            print("Refreshing Google Access Token...")
            try:
                # CHANGED: Using 'GoogleAuthRequest' to avoid name conflict
                creds.refresh(GoogleAuthRequest()) 
                # CHANGED: 'session' is now 'request.session'
                request.session['credentials'] = credentials_to_dict(creds)
            except Exception as e:
                print(f"Failed to refresh token: {e}")
                # CHANGED: 'session' is now 'request.session'
                request.session.clear()
                return None
        else:
            print("Credentials invalid and no refresh token or not expired. Re-authentication needed.")
            # CHANGED: 'session' is now 'request.session'
            request.session.clear()
            return None
    
    try:
        service = build(service_name, version, credentials=creds)
        return service
    except HttpError as error:
        print(f'An error occurred building Google service {service_name} v{version}: {error}')
        return None

def extract_drive_file_id_from_url(url):
    """
    Extracts the Google Drive file ID from various Google Drive/Docs/Sheets/Slides URL formats.
    Returns the file ID string, or None if not found.
    """
    # NO CHANGE: This function is framework-independent.
    match = re.search(r'/d/([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)
    return None


def download_drive_file_content(drive_service, file_id, file_name="unknown"):
    """
    Downloads a Google Drive file's content as plain text.
    
    - If it's a Google Doc/Sheet, it exports as text/plain.
    - If it's a Word document (.docx), it extracts text using python-docx.
    - If it's an Image (JPG, PNG) or PDF, it downloads the bytes 
      and uses the Google Cloud Vision API to extract handwritten text.
    - Otherwise, it attempts a standard text download.
    """
    try:
        file_metadata = drive_service.files().get(fileId=file_id, fields='mimeType, name').execute()
        mime_type = file_metadata.get('mimeType')
        actual_file_name = file_metadata.get('name', file_name)
        print(f"Downloading '{actual_file_name}' (ID: {file_id}) with MIME type: {mime_type}")

        # --- BRANCH 1: Google Workspace Docs ---
        if mime_type.startswith('application/vnd.google-apps'):
            print("File is a Google Doc. Exporting as text/plain.")
            request = drive_service.files().export_media(fileId=file_id, mimeType='text/plain')
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            return fh.getvalue().decode('utf-8')

        # --- BRANCH 2: Word Documents (.docx and .doc) ---
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            print(f"File is a Word document ({mime_type}). Extracting text with python-docx...")
            request = drive_service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            try:
                # Extract text from Word document
                doc = Document(fh)
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Skip empty paragraphs
                        full_text.append(paragraph.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)
                
                extracted_text = '\n'.join(full_text)
                print(f"Successfully extracted {len(extracted_text)} characters from Word document.")
                return extracted_text
            except Exception as docx_error:
                print(f"Error extracting text from Word document: {docx_error}")
                # Fallback: try generic text extraction
                try:
                    fh.seek(0)
                    return fh.getvalue().decode('utf-8')
                except:
                    return None

        # --- BRANCH 3: Images or PDFs ---
        elif mime_type in ['image/jpeg', 'image/png', 'application/pdf']:
            print("File is an Image/PDF. Downloading bytes for Cloud Vision API...")
            request = drive_service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            content_bytes = fh.getvalue()
            
            print("Bytes downloaded. Sending to Cloud Vision API for handwriting detection...")
            
            # --- This is the logic from your detect.py ---
            client = vision.ImageAnnotatorClient()
            image = vision.Image(content=content_bytes)
            image_context = vision.ImageContext(language_hints=["en-t-i0-handwrit"])
            
            response = client.document_text_detection(
                image=image,
                image_context=image_context
            )
            
            if response.error.message:
                print(f"Cloud Vision API Error: {response.error.message}")
                return None
            
            if response.full_text_annotation:
                print("Cloud Vision API successful. Returning extracted text.")
                print(f"--- OCR TEXT FROM {actual_file_name} ---\n{response.full_text_annotation.text}\n---------------------------------")
                return response.full_text_annotation.text
            else:
                print("Cloud Vision API found no text in the image.")
                return "" # Return empty string if no text is found

        # --- BRANCH 4: Other files (e.g., .txt) ---
        else:
            print("File is not a Google Doc, Word document, or Image. Attempting direct media download.")
            request = drive_service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()

            try:
                return fh.getvalue().decode('utf-8')
            except UnicodeDecodeError:
                print(f"UTF-8 decode failed for '{actual_file_name}', trying latin-1.")
                return fh.getvalue().decode('latin-1')

    except HttpError as error:
        print(f'Google Drive API Error downloading file "{actual_file_name}" (ID: {file_id}): {error.resp.status} - {error.content.decode("utf-8")}')
        return None   

    except Exception as e:
        print(f'Unexpected error in download_drive_file_content for "{actual_file_name}" (ID: {file_id}): {e}')
        return None

# --- 3. AUTHENTICATION ROUTES ---

# CHANGED: Converted Flask @app.route to FastAPI @app.get
# ADDED: 'request: Request' parameter for session and url_for
@app.get('/login')
async def login(request: Request):
    """Initiates the Google OAuth 2.0 login flow, redirecting to Google's consent screen."""
    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE,
        scopes=SCOPES, # Use the updated SCOPES list
        # CHANGED: 'url_for' is now 'request.url_for' and requires the endpoint 'name'
        redirect_uri=request.url_for('callback')
    )
    authorization_url, state = flow.authorization_url(access_type='offline', prompt='consent')
    # CHANGED: 'session' is now 'request.session'
    request.session['state'] = state
    # CHANGED: 'redirect' is now 'RedirectResponse'
    return RedirectResponse(authorization_url)

# CHANGED: Converted Flask @app.route to FastAPI @app.get
# ADDED: 'request: Request' parameter
# ADDED: 'name="callback"' to allow request.url_for('callback') to work
@app.get('/auth/callback', name="callback")
async def callback(request: Request):
    """Handles the callback from Google after a user grants or denies permissions."""
    
    # CHANGED: 'session' is now 'request.session'
    # CHANGED: 'request.args.get' is now 'request.query_params.get'
    if 'state' not in request.session or request.session['state'] != request.query_params.get('state'):
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(content={"error": "State mismatch. Possible CSRF attack."}, status_code=400)

    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE,
        scopes=SCOPES, # Use the updated SCOPES list
        state=request.session['state'], # CHANGED: 'session' is now 'request.session'
        redirect_uri=request.url_for('callback') # CHANGED: 'url_for' is now 'request.url_for'
    )
    try:
        flow.fetch_token(authorization_response=str(request.url))
        credentials = flow.credentials
        request.session['credentials'] = credentials_to_dict(credentials)
        request.session.pop('state', None) # CHANGED: 'session' is now 'request.session'
        return RedirectResponse('http://localhost:5173/dashboard')
    except Exception as e:
        print(f"Error fetching token: {e}")
        request.session.clear() # CHANGED: 'session' is now 'request.session'
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(content={"error": "Authentication failed."}, status_code=500)

# CHANGED: Converted Flask @app.route to FastAPI @app.get
# ADDED: 'request: Request' parameter
@app.get('/logout')
async def logout(request: Request):
    """Logs the user out by clearing the session."""
    # CHANGED: 'session' is now 'request.session'
    request.session.clear()
    # CHANGED: 'jsonify' is replaced with returning a dictionary (FastAPI handles it)
    return {"message": "Successfully logged out"}


# --- (Previous code from last snippet) ---
# ... (imports, helpers, and auth routes) ...


# --- 4. API DATA-FETCHING ROUTES ---

# CHANGED: Converted Flask route to FastAPI GET endpoint
# ADDED: 'request: Request' parameter
@app.get('/api/courses')
async def get_courses(request: Request):
    """Fetches and returns a list of the teacher's active Google Classroom courses."""
    # CHANGED: Passed 'request' object to get_google_service
    classroom_service = get_google_service('classroom', 'v1', request)
    if not classroom_service:
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": "User not authenticated or session expired. Please re-login."}, 
            status_code=401
        )
    
    try:
        courses = classroom_service.courses().list(teacherId='me', courseStates=['ACTIVE']).execute()
        # CHANGED: Returned dictionary directly, FastAPI handles 'jsonify'
        return courses.get('courses', [])
    except HttpError as error:
        print(f"Google Classroom API Error in get_courses: {error.resp.status} - {error.content.decode('utf-8')}")
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": f"Failed to fetch courses: {error.content.decode('utf-8')}"}, 
            status_code=error.resp.status
        )

# CHANGED: Converted Flask route to FastAPI GET endpoint
# ADDED: 'request: Request' parameter and type hint for 'course_id'
@app.get('/api/courses/{course_id}/assignments')
async def get_assignments(course_id: str, request: Request):
    """Fetches and returns a list of assignments for a given course ID."""
    # CHANGED: Passed 'request' object to get_google_service
    classroom_service = get_google_service('classroom', 'v1', request)
    if not classroom_service:
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": "User not authenticated or session expired. Please re-login."}, 
            status_code=401
        )
    try:
        all_coursework = classroom_service.courses().courseWork().list(courseId=course_id).execute()
        assignments_only = [
            item for item in all_coursework.get('courseWork', []) 
            if item.get('workType') == 'ASSIGNMENT'
        ]
        # CHANGED: Returned list directly, FastAPI handles 'jsonify'
        return assignments_only

    except HttpError as error:
        print(f"Google Classroom API Error in get_assignments for course {course_id}: {error.resp.status} - {error.content.decode('utf-8')}")
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": f"Failed to fetch assignments: {error.content.decode('utf-8')}"}, 
            status_code=error.resp.status
        )
    except Exception as e:
        print(f"An unexpected error occurred in get_assignments: {e}")
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"}, 
            status_code=500
        )

# CHANGED: Converted Flask route to FastAPI GET endpoint
# ADDED: 'request: Request' parameter and type hints for path parameters
@app.get('/api/courses/{course_id}/assignments/{assignment_id}/submissions')
async def get_submissions(course_id: str, assignment_id: str, request: Request):
    """
    Fetches and returns a list of student submissions for a specific assignment.
    Now includes logic to fetch and attach student names to each submission.
    """
    # CHANGED: Passed 'request' object to get_google_service
    classroom_service = get_google_service('classroom', 'v1', request)
    if not classroom_service:
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": "User not authenticated or session expired. Please re-login."}, 
            status_code=401
        )
    
    try:
        submissions_response = classroom_service.courses().courseWork().studentSubmissions().list(
            courseId=course_id,
            courseWorkId=assignment_id,
            states=['TURNED_IN'] 
        ).execute()
        
        submissions = submissions_response.get('studentSubmissions', [])
        
        # Fetch student profiles for names and enrich submissions
        processed_submissions = []
        for submission in submissions:
            user_id = submission['userId']
            student_name = f"Unknown Student ({user_id})" # Default name in case of error

            try:
                # Need the 'classroom.profile.emails' scope for userProfiles().get
                student_profile = classroom_service.userProfiles().get(userId=user_id).execute()
                student_name = student_profile.get('name', {}).get('fullName', student_name)
            except HttpError as e:
                print(f"Error fetching profile for user {user_id}: {e.resp.status} - {e.content.decode('utf-8')}")
                # If there's an HTTP error, keep the default name
            except Exception as e:
                print(f"An unexpected error occurred while fetching profile for user {user_id}: {e}")
                # If any other error, keep the default name
            
            submission['studentName'] = student_name # Add the name to the submission object
            processed_submissions.append(submission)

        # CHANGED: Returned list directly, FastAPI handles 'jsonify'
        return processed_submissions
    except HttpError as error:
        print(f"Google Classroom API Error in get_submissions for course {course_id}, assignment {assignment_id}: {error.resp.status} - {error.content.decode('utf-8')}")
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": f"Failed to fetch submissions: {error.content.decode('utf-8')}"}, 
            status_code=error.resp.status
        )

import uvicorn # ADDED: For running the FastAPI server
import datetime # ADDED: This import was used in the grading logic

# --- (Previous code from last snippet) ---
# ... (imports, helpers, auth routes, and data-fetching routes) ...


# --- 5. CORE GRADING ROUTE ---

# CHANGED: Converted Flask route to FastAPI POST endpoint
# ADDED: 'request: Request' parameter
@app.post('/api/grade')
async def grade_submission(request: Request):
    global graded_assignments_history 

    # CHANGED: 'request.json' is now 'await request.json()'
    data = await request.json()
    course_id = data.get('course_id')
    course_name = data.get('course_name') 
    assignment_id = data.get('assignment_id')
    assignment_title = data.get('assignment_title') 
    submission_id = data.get('submission_id')
    student_name = data.get('student_name') 
    answer_key_url = data.get('answer_key_url')
    answer_key_text = data.get('answer_key_text')  # NEW: Accept answer key as text
    use_hybrid = data.get('use_hybrid', True)  # NEW: Flag to enable/disable hybrid grading (default: True)

    # Debug logging
    print(f"📥 Received grading request:")
    print(f"   course_id: {course_id}")
    print(f"   course_name: {course_name}")
    print(f"   assignment_id: {assignment_id}")
    print(f"   assignment_title: {assignment_title}")
    print(f"   submission_id: {submission_id}")
    print(f"   student_name: {student_name}")
    print(f"   answer_key_url: {'Present' if answer_key_url else 'Missing'}")
    print(f"   answer_key_text: {'Present' if answer_key_text else 'Missing'}")
    print(f"   use_hybrid: {use_hybrid}")

    if not all([course_id, assignment_id, submission_id, course_name, assignment_title, student_name]):
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        missing_fields = []
        if not course_id: missing_fields.append('course_id')
        if not course_name: missing_fields.append('course_name')
        if not assignment_id: missing_fields.append('assignment_id')
        if not assignment_title: missing_fields.append('assignment_title')
        if not submission_id: missing_fields.append('submission_id')
        if not student_name: missing_fields.append('student_name')
        
        return JSONResponse(
            content={"error": f"Missing required fields: {', '.join(missing_fields)}"},
            status_code=400
        )
    
    # Must have either answer_key_url OR answer_key_text
    if not answer_key_url and not answer_key_text:
        return JSONResponse(
            content={"error": "Must provide either answer_key_url or answer_key_text."},
            status_code=400
        )

    # CHANGED: Passed 'request' object to get_google_service
    classroom_service = get_google_service('classroom', 'v1', request)
    drive_service = get_google_service('drive', 'v3', request)

    if not classroom_service or not drive_service:
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": "User not authenticated or missing Drive/Classroom permissions. Please re-login."},
            status_code=401
        )
    
    try:
        # Get answer key - either from URL or use provided text
        if answer_key_text:
            # Use the provided answer key text directly
            answer_key_content = answer_key_text
            print("Using provided answer key text (generated by AI)")
        else:
            # Download from URL (original behavior)
            print(f"Attempting to extract file ID from answer key URL: {answer_key_url}")
            answer_key_file_id = extract_drive_file_id_from_url(answer_key_url)
            if not answer_key_file_id:
                # CHANGED: 'jsonify' is replaced with 'JSONResponse'
                print(f"❌ Failed to extract file ID from URL: {answer_key_url}")
                return JSONResponse(
                    content={"error": "Invalid Google Drive URL provided for the Answer Key. Please check the URL format (should contain /d/FILE_ID/)."},
                    status_code=400
                )
            
            print(f"✅ Extracted file ID: {answer_key_file_id}")
            print(f"Attempting to download answer key from Google Drive...")
            
            answer_key_content = download_drive_file_content(drive_service, answer_key_file_id, "Answer Key")
            if not answer_key_content:
                print(f"❌ Failed to download answer key. File ID: {answer_key_file_id}")
                return JSONResponse(
                    content={"error": "Failed to download answer key from URL. Please ensure: 1) The file is shared with your Google account, 2) You have View access, 3) The URL is correct."},
                    status_code=500
                )

        assignment_details = classroom_service.courses().courseWork().get(courseId=course_id, id=assignment_id).execute()
        materials = assignment_details.get('materials', [])
        questionnaire_file_id = None
        
        for material in materials:
            if 'driveFile' in material and 'driveFile' in material['driveFile']:
                questionnaire_file_id = material['driveFile']['driveFile']['id']
                print(f"Identified questionnaire file ID: {questionnaire_file_id} from assignment materials.")
                break 
        
        if not questionnaire_file_id:
            # CHANGED: 'jsonify' is replaced with 'JSONResponse'
            return JSONResponse(
                content={"error": "No Google Drive document (questionnaire) found attached to this assignment."},
                status_code=404
            )

        submission_details = classroom_service.courses().courseWork().studentSubmissions().get(
            courseId=course_id, courseWorkId=assignment_id, id=submission_id).execute()
        
        attachments = submission_details.get('assignmentSubmission', {}).get('attachments', [])
        if not attachments:
            # CHANGED: 'jsonify' is replaced with 'JSONResponse'
            return JSONResponse(
                content={"error": "This student has not attached any file to their submission."},
                status_code=404
            )
        
        student_submission_file_id = attachments[0]['driveFile']['id']
        print(f"Identified student submission file ID: {student_submission_file_id}.")


        print("Initiating document downloads...")
        # NOTE: download_drive_file_content is synchronous.
        # For a fully async app, these should be run in a thread pool,
        # but leaving as-is for direct conversion.
        questionnaire_text = download_drive_file_content(drive_service, questionnaire_file_id, "Questionnaire")
        student_submission_text = download_drive_file_content(drive_service, student_submission_file_id, "Student Submission")

        if not all([questionnaire_text, answer_key_content, student_submission_text]):
            # CHANGED: 'jsonify' is replaced with 'JSONResponse'
            return JSONResponse(
                content={"error": "Failed to download text content from one or more required documents. Check file permissions or existence."},
                status_code=500
            )
        
        print("Documents downloaded. Constructing Gemini prompt and calling AI...")
        model = genai.GenerativeModel(model_name="gemini-2.5-flash") # Using 1.5 Flash as a modern equivalent
        
        prompt = f"""
        You are an expert AI teaching assistant for a Google Classroom assignment titled "{assignment_details.get('title', 'Unknown Assignment')}". Your task is to rigorously grade the student's submission, providing a score out of 100, and comprehensive feedback.

        --- QUESTIONNAIRE (The questions/tasks presented to the student) ---
        {questionnaire_text}

        --- OFFICIAL ANSWER KEY (The expected correct responses/solutions) ---
        {answer_key_content}

        --- STUDENT'S SUBMISSION (The student's actual answers/work) ---
        {student_submission_text}

        --- GRADING INSTRUCTIONS ---
        1.  **Understanding the Task:** Carefully read the QUESTIONNAIRE to grasp the specific requirements and learning objectives.

        2.  **Content Accuracy & Completeness:** Compare the STUDENT'S SUBMISSION against the OFFICIAL ANSWER KEY.
            * How accurately does the student address each question/task?
            * Is the information presented correct?
            * Are all parts of the question/task attempted and completed?
            * **Award points for partially correct or reasonable attempts. Avoid giving a 0 unless the submission is entirely blank, off-topic, or completely nonsensical.** Even minimal effort to address the prompt should receive some credit.

        3.  **Structure & Clarity:** Evaluate the organization, clarity, and readability of the student's response.

        4.  **Meaning & Comprehension:** Assess the student's understanding of the concepts. Does their submission demonstrate comprehension, or is it just rote memorization/copying?

        5.  **Assign a Numerical Grade (0-100):** Based on the above criteria, assign a numerical grade, using the following guidelines to achieve scores between 70-80 for conceptually correct but less precise answers:
            * **90-100 (Excellent):** Answers are accurate, complete, well-structured, and demonstrate deep comprehension. Critically, they are also *precise* and leverage key terminology from the answer key where appropriate.
            * **75-89 (Good/Strong):** Answers are *conceptually correct* and show good understanding, but might lack the highest level of precision or miss some specific key terminology from the answer key. They are clear, mostly complete, but could be more refined.
            * **50-74 (Fair/Developing):** Answers are partially correct, contain some inaccuracies, or are vague. They may demonstrate some understanding but require significant improvement in content, clarity, or completeness.
            * **< 50 (Limited/Poor):** Answers are largely incorrect, off-topic, or show minimal understanding. This category should only be used if attempts are very weak or absent.

        6.  **Provide Comprehensive Feedback:**
            * Start with positive aspects or areas where the student demonstrated understanding.
            * Clearly explain where points were lost, referencing specific parts of the questionnaire or answer key.
            * For "Good/Strong" answers, specifically suggest how they could make their explanation more precise or complete by integrating relevant key terms or more detailed examples.
            * Suggest concrete steps for improvement.

        7.  **Justify the Grade (Briefly):** Include a short sentence explaining the overall reasoning for the assigned score, linking it to the guidelines above (e.g., "Conceptually solid but lacked some precision and specific terminology for full marks.")

        8.  **Format your response STRICTLY as follows, with no extra text before or after, ensuring clear separation for parsing:
        GRADE: [SCORE]/100
        GRADE_JUSTIFICATION: [A brief, one-sentence reason for the score]
        FEEDBACK: [Your detailed feedback paragraph here, covering all points from instruction 6]
        """
        
        # CHANGED: Switched to the async version of the call
        response = await model.generate_content_async(prompt)
        
        print("Gemini response received. Parsing...")
        grade_text_output = response.text
        
        grade_match = re.search(r'GRADE:\s*(\d+)/100', grade_text_output, re.IGNORECASE)
        justification_match = re.search(r'GRADE_JUSTIFICATION:\s*(.*)', grade_text_output, re.IGNORECASE)
        feedback_match = re.search(r'FEEDBACK:\s*(.*)', grade_text_output, re.IGNORECASE | re.DOTALL)

        if not grade_match or not feedback_match or not justification_match:
            print(f"Gemini response was not in the expected format. Raw response:\n{grade_text_output}")
            # CHANGED: 'jsonify' is replaced with 'JSONResponse'
            return JSONResponse(
                content={"error": "Failed to parse grade, justification, or feedback from AI response. Please ensure Gemini's output adheres to the specified format."},
                status_code=500
            )
            
        ai_grade = int(grade_match.group(1))
        ai_justification = justification_match.group(1).strip()
        ai_feedback = feedback_match.group(1).strip()
        
        # --- CONDITIONAL HYBRID GRADING ---
        if use_hybrid:
            # Use hybrid grading (MiniLM + Gemini) for Grade WITH Key workflow
            print("=" * 60)
            print("Starting hybrid grading (MiniLM + Gemini)...")
            print("=" * 60)
            
            # Get Gemini's grade
            gemini_grade = ai_grade
            print(f"📝 Gemini grade: {gemini_grade}/100")
            
            # Try to get MiniLM semantic similarity score
            print(f"🔍 Calculating MiniLM semantic similarity...")
            minilm_similarity = get_minilm_semantic_score(answer_key_content, student_submission_text)
            
            if minilm_similarity is not None:
                print(f"✅ MiniLM similarity score: {minilm_similarity:.4f}")
                minilm_grade = normalize_minilm_score_to_grade(minilm_similarity)
                print(f"📊 MiniLM normalized grade: {minilm_grade}/100")
            else:
                print(f"⚠️ MiniLM model not available, using Gemini only")
                minilm_grade = None
            
            # Compare and decide final grade
            grade_analysis = compare_minilm_and_gemini_grades(minilm_grade, gemini_grade)
            
            final_grade = grade_analysis['final_grade']
            confidence_level = grade_analysis['confidence']
            grading_method = grade_analysis['method']
            
            print(f"🎯 Final grade: {final_grade}/100")
            print(f"📊 Confidence: {confidence_level}")
            print(f"🔧 Method: {grading_method}")
            print(f"💡 {grade_analysis['recommendation']}")
            print("=" * 60)
            
            # Build comprehensive justification
            if minilm_grade is not None:
                grade_justification = (
                    f"{ai_justification} "
                    f"[Verified with MiniLM semantic model: {minilm_similarity:.3f} similarity, "
                    f"{minilm_grade}/100. Difference: {grade_analysis['difference']}pts]"
                )
                remarks = (
                    f"Hybrid grading ({grading_method}): MiniLM similarity={minilm_similarity:.3f}, "
                    f"MiniLM grade={minilm_grade}, Gemini grade={gemini_grade}, Final={final_grade}. "
                    f"{grade_analysis['recommendation']}"
                )
            else:
                grade_justification = f"{ai_justification} [Gemini-only grading]"
                remarks = "Graded using Gemini AI only (MiniLM model not available)."
            
            feedback_str = ai_feedback
        else:
            # Use Gemini-only grading for Grade WITHOUT Key workflow
            print("=" * 60)
            print("Using Gemini-only grading (hybrid disabled)...")
            print("=" * 60)
            
            final_grade = ai_grade
            confidence_level = "high"
            grading_method = "gemini_only"
            grade_justification = f"{ai_justification} [Gemini-only grading]"
            remarks = "Graded using Gemini AI only (no hybrid model)."
            feedback_str = ai_feedback
            minilm_grade = None
            grade_analysis = None
            
            print(f"🎯 Final grade: {final_grade}/100")
            print(f"📝 Method: Gemini only")
            print("=" * 60)
        
        # --- END CONDITIONAL GRADING ---
        
        # import datetime # This was here, moved to top
        graded_item = {
            "course_id": course_id,
            "course_name": course_name,
            "assignment_id": assignment_id,
            "assignment_title": assignment_title,
            "submission_id": submission_id,
            "student_name": student_name,
            "assignedGrade": final_grade,
            "confidence": confidence_level,
            "grading_method": grading_method,
            "feedback": feedback_str,
            "grade_justification": grade_justification,
            "remarks": remarks,
            "timestamp": datetime.datetime.now().isoformat()
        }
        
        # Save to MongoDB (with fallback to in-memory)
        if grades_collection is not None:
            try:
                result = grades_collection.insert_one(graded_item.copy())
                print(f"✅ Grade saved to MongoDB with ID: {result.inserted_id}")
                
                # Update student's record
                students_collection.update_one(
                    {"student_name": student_name, "course_id": course_id},
                    {
                        "$set": {
                            "student_name": student_name,
                            "course_id": course_id,
                            "course_name": course_name,
                            "last_updated": datetime.datetime.now().isoformat()
                        },
                        "$inc": {"total_assignments": 1},
                        "$push": {
                            "grades_history": {
                                "assignment_id": assignment_id,
                                "assignment_title": assignment_title,
                                "grade": final_grade,
                                "timestamp": graded_item["timestamp"]
                            }
                        }
                    },
                    upsert=True
                )
                print(f"✅ Student profile updated for {student_name}")
            except Exception as mongo_error:
                print(f"⚠️ MongoDB save error: {mongo_error}")
                print("⚠️ Falling back to in-memory storage")
                graded_assignments_history.append(graded_item)
        else:
            # Fallback to in-memory if MongoDB not connected
            graded_assignments_history.append(graded_item)
            print("⚠️ Using in-memory storage (MongoDB not connected)")

        print(f"Gemini Grade: {final_grade}/100. Providing review for dashboard display only (no Classroom update).")
        
        # Build response with grading information
        response_data = {
            "message": f"{'Hybrid' if use_hybrid else 'Gemini-only'} grading complete. Review provided for dashboard display only.",
            "assignedGrade": final_grade,
            "confidence": confidence_level,
            "grading_method": grading_method,
            "feedback": feedback_str,
            "grade_justification": grade_justification,
            "remarks": remarks,
            "status": "review_only", 
            "graded_history": graded_assignments_history
        }
        
        # Add separate grades when hybrid grading is enabled and available
        if use_hybrid and minilm_grade is not None and grade_analysis is not None:
            response_data["minilm_grade"] = minilm_grade
            response_data["gemini_grade"] = ai_grade
            response_data["grade_difference"] = grade_analysis.get('difference')
            response_data["minilm_similarity"] = round(get_minilm_semantic_score(answer_key_content, student_submission_text), 4)
            
            # If low/medium confidence, flag for detailed review
            if confidence_level in ['low', 'medium'] and grade_analysis.get('difference', 0) > 15:
                response_data["needs_review"] = True
                response_data["review_reason"] = f"MiniLM and Gemini disagree by {grade_analysis['difference']} points"
        
        # CHANGED: 'jsonify' is replaced with returning a dictionary
        return response_data


    except HttpError as error:
        error_details = error.content.decode('utf-8')
        print(f"Google API Error in grading: {error.resp.status} - {error_details}")
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": f"Google API Error: {error_details}"},
            status_code=error.resp.status
        )
    except Exception as e:
        print(f"An unexpected error occurred during grading: {e}")
        # CHANGED: 'jsonify' is replaced with 'JSONResponse'
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"},
            status_code=500
        )

# --- New route to fetch the entire grading history ---

# CHANGED: Converted Flask route to FastAPI GET endpoint
@app.get('/api/graded_history')
async def get_graded_history(course_id: str = None, assignment_id: str = None):
    """
    Returns the list of all graded assignments from MongoDB.
    Can filter by course_id and/or assignment_id.
    """
    try:
        # Build query filter
        query = {}
        if course_id:
            query['course_id'] = course_id
        if assignment_id:
            query['assignment_id'] = assignment_id
        
        # Fetch from MongoDB if available
        if grades_collection is not None:
            grades = list(grades_collection.find(
                query,
                {'_id': 0}  # Exclude MongoDB's _id field
            ).sort('timestamp', -1))  # Most recent first
            print(f"📊 Fetched {len(grades)} grades from MongoDB")
            return grades
        else:
            # Fallback to in-memory storage
            print("⚠️ Using in-memory storage (MongoDB not connected)")
            if course_id or assignment_id:
                # Apply filters manually for in-memory
                filtered = [
                    g for g in graded_assignments_history
                    if (not course_id or g.get('course_id') == course_id) and
                       (not assignment_id or g.get('assignment_id') == assignment_id)
                ]
                return filtered
            return graded_assignments_history
    except Exception as e:
        print(f"❌ Error fetching graded history: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


# --- 7. GRADING HUB ENDPOINTS (New Feature) ---

@app.post('/api/grade-with-model')
async def grade_with_model(request: Request):
    """
    Path 1: Grade with Provided Key (Custom ML Model Workflow)
    This endpoint uses a custom ML model to grade student submissions.
    """
    data = await request.json()
    course_id = data.get('course_id')
    assignment_id = data.get('assignment_id')
    submission_id = data.get('submission_id')
    answer_key_url = data.get('answer_key_url')
    student_name = data.get('student_name')

    # Debug logging
    print(f"📥 Received /api/grade-with-model request:")
    print(f"   course_id: {course_id}")
    print(f"   assignment_id: {assignment_id}")
    print(f"   submission_id: {submission_id}")
    print(f"   answer_key_url: {answer_key_url}")
    print(f"   student_name: {student_name}")

    if not all([course_id, assignment_id, submission_id, answer_key_url]):
        missing_fields = []
        if not course_id: missing_fields.append('course_id')
        if not assignment_id: missing_fields.append('assignment_id')
        if not submission_id: missing_fields.append('submission_id')
        if not answer_key_url: missing_fields.append('answer_key_url')
        
        return JSONResponse(
            content={"error": f"Missing required fields: {', '.join(missing_fields)}"},
            status_code=400
        )

    drive_service = get_google_service('drive', 'v3', request)
    classroom_service = get_google_service('classroom', 'v1', request)

    if not drive_service or not classroom_service:
        return JSONResponse(
            content={"error": "User not authenticated. Please re-login."},
            status_code=401
        )

    try:
        # Extract answer key file ID from URL
        answer_key_file_id = extract_drive_file_id_from_url(answer_key_url)
        if not answer_key_file_id:
            return JSONResponse(
                content={"error": "Invalid Google Drive URL provided for the Answer Key."},
                status_code=400
            )

        # Get student submission file ID
        submission_details = classroom_service.courses().courseWork().studentSubmissions().get(
            courseId=course_id, courseWorkId=assignment_id, id=submission_id).execute()
        
        attachments = submission_details.get('assignmentSubmission', {}).get('attachments', [])
        if not attachments:
            return JSONResponse(
                content={"error": "This student has not attached any file to their submission."},
                status_code=404
            )
        
        student_submission_file_id = attachments[0]['driveFile']['id']

        # Download both files with OCR support
        print("Downloading answer key and student submission...")
        answer_key_text = download_drive_file_content(drive_service, answer_key_file_id, "Answer Key")
        student_submission_text = download_drive_file_content(drive_service, student_submission_file_id, "Student Submission")

        if not answer_key_text or not student_submission_text:
            return JSONResponse(
                content={"error": "Failed to download one or more required documents."},
                status_code=500
            )

        # TODO: Replace this placeholder with actual ML model call
        # For now, using a simple placeholder response
        print("Feeding to custom ML model (placeholder)...")
        
        # Placeholder ML model response
        ml_grade = 85  # This would come from your ML model
        ml_feedback = "Custom ML model grading completed. This is a placeholder response."
        
        return {
            "assignedGrade": ml_grade,
            "feedback": ml_feedback,
            "status": "graded_with_model"
        }

    except HttpError as error:
        error_details = error.content.decode('utf-8')
        print(f"Google API Error: {error.resp.status} - {error_details}")
        return JSONResponse(
            content={"error": f"Google API Error: {error_details}"},
            status_code=error.resp.status
        )
    except Exception as e:
        print(f"Unexpected error in grade_with_model: {e}")
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"},
            status_code=500
        )


@app.post('/api/generate-initial-key')
async def generate_initial_key(request: Request):
    """
    Path 2 - Step 1: Generate initial answer key using Gemini
    Extracts questionnaire from assignment materials automatically
    """
    data = await request.json()
    course_id = data.get('course_id')
    assignment_id = data.get('assignment_id')

    if not all([course_id, assignment_id]):
        return JSONResponse(
            content={"error": "Missing required data: course_id or assignment_id."},
            status_code=400
        )

    classroom_service = get_google_service('classroom', 'v1', request)
    drive_service = get_google_service('drive', 'v3', request)

    if not classroom_service or not drive_service:
        return JSONResponse(
            content={"error": "User not authenticated. Please re-login."},
            status_code=401
        )

    try:
        # Get assignment details to find questionnaire (same as in /api/grade)
        print(f"Fetching assignment details for course {course_id}, assignment {assignment_id}...")
        assignment_details = classroom_service.courses().courseWork().get(
            courseId=course_id, id=assignment_id).execute()
        
        materials = assignment_details.get('materials', [])
        questionnaire_file_id = None
        
        for material in materials:
            if 'driveFile' in material and 'driveFile' in material['driveFile']:
                questionnaire_file_id = material['driveFile']['driveFile']['id']
                print(f"Found questionnaire file ID: {questionnaire_file_id}")
                break
        
        if not questionnaire_file_id:
            return JSONResponse(
                content={"error": "No questionnaire document found attached to this assignment. Please ensure a Google Drive document is attached as assignment material."},
                status_code=404
            )

        # Download questionnaire with OCR support
        print("Downloading questionnaire with OCR support...")
        questionnaire_text = download_drive_file_content(drive_service, questionnaire_file_id, "Questionnaire")

        if not questionnaire_text:
            return JSONResponse(
                content={"error": "Failed to download questionnaire document."},
                status_code=500
            )

        print(f"Questionnaire extracted successfully. Text length: {len(questionnaire_text)} characters")

        # Generate answer key using Gemini
        print("Generating initial answer key with Gemini AI...")
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        
        prompt = f"""You are an expert teacher. Generate a high-quality, detailed answer key for the following questions.

--- QUESTIONS ---
{questionnaire_text}

--- INSTRUCTIONS ---
1. Provide clear, comprehensive answers to each question
2. Include key concepts and terminology that should be present in student responses
3. Structure your answers in a way that makes it easy to grade against
4. Be thorough but concise
5. Format the answer key in a structured way (e.g., "Question 1: [answer]", "Question 2: [answer]", etc.)

Please provide the answer key in a clear, organized format."""

        response = await model.generate_content_async(prompt)
        answer_key = response.text

        print("Answer key generated successfully!")
        return {
            "answer_key": answer_key,
            "status": "success"
        }

    except HttpError as error:
        error_details = error.content.decode('utf-8')
        print(f"Google API Error: {error.resp.status} - {error_details}")
        return JSONResponse(
            content={"error": f"Google API Error: {error_details}"},
            status_code=error.resp.status
        )
    except Exception as e:
        print(f"Unexpected error in generate_initial_key: {e}")
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"},
            status_code=500
        )


@app.post('/api/refine-key')
async def refine_key(request: Request):
    """
    Path 2 - Step 2: Refine the answer key based on teacher feedback
    """
    data = await request.json()
    course_id = data.get('course_id')
    assignment_id = data.get('assignment_id')
    current_key = data.get('current_key')
    feedback = data.get('feedback')

    if not all([course_id, assignment_id, current_key, feedback]):
        return JSONResponse(
            content={"error": "Missing required data: course_id, assignment_id, current_key, or feedback."},
            status_code=400
        )

    try:
        print("Refining answer key with Gemini based on teacher feedback...")
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        
        prompt = f"""You are an AI assistant helping a teacher refine an answer key for their assignment.

--- CURRENT ANSWER KEY ---
{current_key}

--- TEACHER'S FEEDBACK FOR IMPROVEMENT ---
{feedback}

--- INSTRUCTIONS ---
1. Carefully read the teacher's feedback and understand their concerns or suggestions
2. Generate an improved version of the answer key that addresses all feedback points
3. Maintain the overall structure and format of the original answer key
4. Make the refinements clear and well-integrated
5. Ensure the refined answer key is comprehensive and ready for grading

Please provide the complete refined answer key below:"""

        response = await model.generate_content_async(prompt)
        refined_key = response.text

        print("Answer key refined successfully!")
        return {
            "refined_key": refined_key,
            "status": "success"
        }

    except Exception as e:
        print(f"Unexpected error in refine_key: {e}")
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"},
            status_code=500
        )


@app.post('/api/grade-with-gemini')
async def grade_with_gemini(request: Request):
    """
    Path 2 - Step 3: Grade ALL student submissions using Gemini with the approved answer key
    """
    global graded_assignments_history

    data = await request.json()
    course_id = data.get('course_id')
    assignment_id = data.get('assignment_id')
    approved_key = data.get('approved_key')

    if not all([course_id, assignment_id, approved_key]):
        return JSONResponse(
            content={"error": "Missing required data: course_id, assignment_id, or approved_key."},
            status_code=400
        )

    classroom_service = get_google_service('classroom', 'v1', request)
    drive_service = get_google_service('drive', 'v3', request)

    if not classroom_service or not drive_service:
        return JSONResponse(
            content={"error": "User not authenticated. Please re-login."},
            status_code=401
        )

    try:
        # Get assignment details for questionnaire and metadata
        print(f"Fetching assignment details for course {course_id}, assignment {assignment_id}...")
        assignment_details = classroom_service.courses().courseWork().get(
            courseId=course_id, id=assignment_id).execute()
        
        assignment_title = assignment_details.get('title', 'Unknown Assignment')
        
        # Get course details for course name
        course_details = classroom_service.courses().get(id=course_id).execute()
        course_name = course_details.get('name', 'Unknown Course')
        
        # Find questionnaire in assignment materials
        materials = assignment_details.get('materials', [])
        questionnaire_file_id = None
        
        for material in materials:
            if 'driveFile' in material and 'driveFile' in material['driveFile']:
                questionnaire_file_id = material['driveFile']['driveFile']['id']
                print(f"Found questionnaire file ID: {questionnaire_file_id}")
                break
        
        if not questionnaire_file_id:
            return JSONResponse(
                content={"error": "No questionnaire document found attached to this assignment."},
                status_code=404
            )

        # Download questionnaire once (will be used for all students)
        print("Downloading questionnaire...")
        questionnaire_text = download_drive_file_content(drive_service, questionnaire_file_id, "Questionnaire")

        if not questionnaire_text:
            return JSONResponse(
                content={"error": "Failed to download questionnaire document."},
                status_code=500
            )

        # Get ALL student submissions
        print("Fetching all student submissions...")
        submissions_response = classroom_service.courses().courseWork().studentSubmissions().list(
            courseId=course_id,
            courseWorkId=assignment_id,
            states=['TURNED_IN']
        ).execute()
        
        submissions = submissions_response.get('studentSubmissions', [])
        
        if not submissions:
            return JSONResponse(
                content={"error": "No student submissions found for this assignment."},
                status_code=404
            )

        print(f"Found {len(submissions)} submissions to grade.")
        
        # Grade each submission
        graded_submissions = []
        graded_count = 0
        
        for submission in submissions:
            try:
                submission_id = submission['id']
                user_id = submission['userId']
                
                # Get student name
                student_name = f"Unknown Student ({user_id})"
                try:
                    student_profile = classroom_service.userProfiles().get(userId=user_id).execute()
                    student_name = student_profile.get('name', {}).get('fullName', student_name)
                except Exception as e:
                    print(f"Could not fetch name for user {user_id}: {e}")
                
                # Check if submission has attachments
                attachments = submission.get('assignmentSubmission', {}).get('attachments', [])
                if not attachments:
                    print(f"Skipping submission {submission_id} - no attachments")
                    graded_submissions.append({
                        "submission_id": submission_id,
                        "student_name": student_name,
                        "status": "skipped",
                        "error": "No file attached"
                    })
                    continue
                
                student_submission_file_id = attachments[0]['driveFile']['id']
                
                # Download student submission
                print(f"Downloading submission for {student_name}...")
                student_submission_text = download_drive_file_content(
                    drive_service, 
                    student_submission_file_id, 
                    f"Submission - {student_name}"
                )

                if not student_submission_text:
                    print(f"Failed to download submission for {student_name}")
                    graded_submissions.append({
                        "submission_id": submission_id,
                        "student_name": student_name,
                        "status": "error",
                        "error": "Failed to download submission"
                    })
                    continue

                # Grade using ONLY Gemini (NO MiniLM/hybrid for this route)
                print(f"Grading submission for {student_name} with Gemini only...")
                model = genai.GenerativeModel(model_name="gemini-2.5-flash")
                
                prompt = f"""
        You are an expert AI teaching assistant for a Google Classroom assignment titled "{assignment_title}". Your task is to rigorously grade the student's submission, providing a score out of 100, and comprehensive feedback.

        --- QUESTIONNAIRE (The questions/tasks presented to the student) ---
        {questionnaire_text}

        --- OFFICIAL ANSWER KEY (The expected correct responses/solutions) ---
        {approved_key}

        --- STUDENT'S SUBMISSION (The student's actual answers/work) ---
        {student_submission_text}

        --- GRADING INSTRUCTIONS ---
        1.  **Understanding the Task:** Carefully read the QUESTIONNAIRE to grasp the specific requirements and learning objectives.

        2.  **Content Accuracy & Completeness:** Compare the STUDENT'S SUBMISSION against the OFFICIAL ANSWER KEY.
            * How accurately does the student address each question/task?
            * Is the information presented correct?
            * Are all parts of the question/task attempted and completed?
            * **Award points for partially correct or reasonable attempts. Avoid giving a 0 unless the submission is entirely blank, off-topic, or completely nonsensical.** Even minimal effort to address the prompt should receive some credit.

        3.  **Structure & Clarity:** Evaluate the organization, clarity, and readability of the student's response.

        4.  **Meaning & Comprehension:** Assess the student's understanding of the concepts. Does their submission demonstrate comprehension, or is it just rote memorization/copying?

        5.  **Assign a Numerical Grade (0-100):** Based on the above criteria, assign a numerical grade, using the following guidelines to achieve scores between 70-80 for conceptually correct but less precise answers:
            * **90-100 (Excellent):** Answers are accurate, complete, well-structured, and demonstrate deep comprehension. Critically, they are also *precise* and leverage key terminology from the answer key where appropriate.
            * **75-89 (Good/Strong):** Answers are *conceptually correct* and show good understanding, but might lack the highest level of precision or miss some specific key terminology from the answer key. They are clear, mostly complete, but could be more refined.
            * **50-74 (Fair/Developing):** Answers are partially correct, contain some inaccuracies, or are vague. They may demonstrate some understanding but require significant improvement in content, clarity, or completeness.
            * **< 50 (Limited/Poor):** Answers are largely incorrect, off-topic, or show minimal understanding. This category should only be used if attempts are very weak or absent.

        6.  **Provide Comprehensive Feedback:**
            * Start with positive aspects or areas where the student demonstrated understanding.
            * Clearly explain where points were lost, referencing specific parts of the questionnaire or answer key.
            * For "Good/Strong" answers, specifically suggest how they could make their explanation more precise or complete by integrating relevant key terms or more detailed examples.
            * Suggest concrete steps for improvement.

        7.  **Justify the Grade (Briefly):** Include a short sentence explaining the overall reasoning for the assigned score, linking it to the guidelines above (e.g., "Conceptually solid but lacked some precision and specific terminology for full marks.")

        8.  **Format your response STRICTLY as follows, with no extra text before or after, ensuring clear separation for parsing:
        GRADE: [SCORE]/100
        GRADE_JUSTIFICATION: [A brief, one-sentence reason for the score]
        FEEDBACK: [Your detailed feedback paragraph here, covering all points from instruction 6]
        """
                
                response = await model.generate_content_async(prompt)
                grade_text_output = response.text
                
                # Parse the response
                grade_match = re.search(r'GRADE:\s*(\d+)/100', grade_text_output, re.IGNORECASE)
                justification_match = re.search(r'GRADE_JUSTIFICATION:\s*(.*)', grade_text_output, re.IGNORECASE)
                feedback_match = re.search(r'FEEDBACK:\s*(.*)', grade_text_output, re.IGNORECASE | re.DOTALL)

                if not grade_match or not feedback_match or not justification_match:
                    print(f"Failed to parse Gemini response for {student_name}")
                    graded_submissions.append({
                        "submission_id": submission_id,
                        "student_name": student_name,
                        "status": "error",
                        "error": "Failed to parse AI response"
                    })
                    continue
                    
                final_grade = int(grade_match.group(1))
                grade_justification = justification_match.group(1).strip()
                feedback_str = feedback_match.group(1).strip()
                
                # Store graded item
                graded_item = {
                    "course_id": course_id,
                    "course_name": course_name,
                    "assignment_id": assignment_id,
                    "assignment_title": assignment_title,
                    "submission_id": submission_id,
                    "student_name": student_name,
                    "assignedGrade": final_grade,
                    "confidence": "high",
                    "grading_method": "gemini_only",
                    "feedback": feedback_str,
                    "grade_justification": grade_justification,
                    "remarks": "Graded using Gemini AI only (no hybrid model)",
                    "timestamp": datetime.datetime.now().isoformat()
                }
                
                # Save to MongoDB (with fallback to in-memory)
                if grades_collection is not None:
                    try:
                        result = grades_collection.insert_one(graded_item.copy())
                        print(f"✅ Grade saved to MongoDB for {student_name}")
                        
                        # Update student's record
                        students_collection.update_one(
                            {"student_name": student_name, "course_id": course_id},
                            {
                                "$set": {
                                    "student_name": student_name,
                                    "course_id": course_id,
                                    "course_name": course_name,
                                    "last_updated": datetime.datetime.now().isoformat()
                                },
                                "$inc": {"total_assignments": 1},
                                "$push": {
                                    "grades_history": {
                                        "assignment_id": assignment_id,
                                        "assignment_title": assignment_title,
                                        "grade": final_grade,
                                        "timestamp": graded_item["timestamp"]
                                    }
                                }
                            },
                            upsert=True
                        )
                    except Exception as mongo_error:
                        print(f"⚠️ MongoDB save error for {student_name}: {mongo_error}")
                        graded_assignments_history.append(graded_item)
                else:
                    graded_assignments_history.append(graded_item)
                
                graded_submissions.append({
                    "submission_id": submission_id,
                    "student_name": student_name,
                    "assignedGrade": final_grade,
                    "feedback": feedback_str,
                    "grade_justification": grade_justification,
                    "status": "success"
                })
                
                graded_count += 1
                print(f"Successfully graded {student_name}: {final_grade}/100")
                
            except Exception as e:
                print(f"Error grading submission {submission.get('id', 'unknown')}: {e}")
                graded_submissions.append({
                    "submission_id": submission.get('id', 'unknown'),
                    "student_name": student_name if 'student_name' in locals() else "Unknown",
                    "status": "error",
                    "error": str(e)
                })

        print(f"Grading complete! Successfully graded {graded_count} out of {len(submissions)} submissions.")
        
        return {
            "graded_count": graded_count,
            "total_submissions": len(submissions),
            "submissions": graded_submissions,
            "status": "complete"
        }

    except HttpError as error:
        error_details = error.content.decode('utf-8')
        print(f"Google API Error: {error.resp.status} - {error_details}")
        return JSONResponse(
            content={"error": f"Google API Error: {error_details}"},
            status_code=error.resp.status
        )
    except Exception as e:
        print(f"Unexpected error in grade_with_gemini: {e}")
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"},
            status_code=500
        )


@app.post('/api/export-grades-to-sheet')
async def export_grades_to_sheet(request: Request):
    """
    Export graded submissions to a Google Sheet and save it to the teacher's Drive.
    
    Expected data format:
    {
        "course_name": "Machine Learning 101",
        "assignment_title": "Midterm Exam",
        "graded_submissions": [
            {
                "student_name": "John Doe",
                "assignedGrade": 85,
                "feedback": "Good work on questions 1-3..."
            },
            ...
        ]
    }
    """
    data = await request.json()
    course_name = data.get('course_name', 'Unknown Course')
    assignment_title = data.get('assignment_title', 'Unknown Assignment')
    graded_submissions = data.get('graded_submissions', [])

    if not graded_submissions:
        return JSONResponse(
            content={"error": "No graded submissions provided for export."},
            status_code=400
        )

    sheets_service = get_google_service('sheets', 'v4', request)
    drive_service = get_google_service('drive', 'v3', request)

    if not sheets_service or not drive_service:
        return JSONResponse(
            content={"error": "User not authenticated. Please re-login."},
            status_code=401
        )

    try:
        # Create a new Google Sheet
        spreadsheet_title = f"{course_name} - {assignment_title} - Graded Results"
        print(f"Creating Google Sheet: {spreadsheet_title}")
        
        spreadsheet_body = {
            'properties': {
                'title': spreadsheet_title
            },
            'sheets': [{
                'properties': {
                    'title': 'Grades',
                    'gridProperties': {
                        'frozenRowCount': 1  # Freeze header row
                    }
                }
            }]
        }
        
        spreadsheet = sheets_service.spreadsheets().create(
            body=spreadsheet_body,
            fields='spreadsheetId,spreadsheetUrl,sheets'
        ).execute()
        
        spreadsheet_id = spreadsheet.get('spreadsheetId')
        spreadsheet_url = spreadsheet.get('spreadsheetUrl')
        
        # Get the actual sheet ID from the created spreadsheet
        sheet_id = spreadsheet['sheets'][0]['properties']['sheetId']
        
        print(f"Google Sheet created with ID: {spreadsheet_id}, Sheet ID: {sheet_id}")
        
        # Prepare data for the sheet
        # Header row
        values = [
            ['Student Name', 'Grade', 'Out Of', 'Feedback/Review']
        ]
        
        # Data rows
        for submission in graded_submissions:
            student_name = submission.get('student_name', 'Unknown Student')
            grade = submission.get('assignedGrade', 0)
            feedback = submission.get('feedback', 'No feedback provided')
            
            values.append([
                student_name,
                grade,
                100,  # Out of 100
                feedback
            ])
        
        # Update the sheet with data
        print(f"Adding {len(values)-1} student records to sheet...")
        range_name = 'Grades!A1'
        value_input_option = 'USER_ENTERED'
        
        body = {
            'values': values
        }
        
        sheets_service.spreadsheets().values().update(
            spreadsheetId=spreadsheet_id,
            range=range_name,
            valueInputOption=value_input_option,
            body=body
        ).execute()
        
        # Format the sheet (bold headers, borders, auto-resize columns)
        print("Applying formatting to sheet...")
        requests = [
            # Bold header row
            {
                'repeatCell': {
                    'range': {
                        'sheetId': sheet_id,
                        'startRowIndex': 0,
                        'endRowIndex': 1
                    },
                    'cell': {
                        'userEnteredFormat': {
                            'textFormat': {
                                'bold': True,
                                'fontSize': 11
                            },
                            'backgroundColor': {
                                'red': 0.9,
                                'green': 0.9,
                                'blue': 0.9
                            },
                            'horizontalAlignment': 'CENTER'
                        }
                    },
                    'fields': 'userEnteredFormat(textFormat,backgroundColor,horizontalAlignment)'
                }
            },
            # Add borders to all cells
            {
                'updateBorders': {
                    'range': {
                        'sheetId': sheet_id,
                        'startRowIndex': 0,
                        'endRowIndex': len(values),
                        'startColumnIndex': 0,
                        'endColumnIndex': 4
                    },
                    'top': {'style': 'SOLID', 'width': 1},
                    'bottom': {'style': 'SOLID', 'width': 1},
                    'left': {'style': 'SOLID', 'width': 1},
                    'right': {'style': 'SOLID', 'width': 1},
                    'innerHorizontal': {'style': 'SOLID', 'width': 1},
                    'innerVertical': {'style': 'SOLID', 'width': 1}
                }
            },
            # Auto-resize columns
            {
                'autoResizeDimensions': {
                    'dimensions': {
                        'sheetId': sheet_id,
                        'dimension': 'COLUMNS',
                        'startIndex': 0,
                        'endIndex': 4
                    }
                }
            },
            # Center-align Grade and Out Of columns
            {
                'repeatCell': {
                    'range': {
                        'sheetId': sheet_id,
                        'startRowIndex': 1,
                        'endRowIndex': len(values),
                        'startColumnIndex': 1,
                        'endColumnIndex': 3
                    },
                    'cell': {
                        'userEnteredFormat': {
                            'horizontalAlignment': 'CENTER'
                        }
                    },
                    'fields': 'userEnteredFormat.horizontalAlignment'
                }
            }
        ]
        
        sheets_service.spreadsheets().batchUpdate(
            spreadsheetId=spreadsheet_id,
            body={'requests': requests}
        ).execute()
        
        print(f"Sheet formatted successfully! URL: {spreadsheet_url}")
        
        return {
            "spreadsheet_id": spreadsheet_id,
            "spreadsheet_url": spreadsheet_url,
            "status": "success",
            "message": f"Successfully exported {len(graded_submissions)} graded submissions to Google Sheets",
            "student_count": len(graded_submissions)
        }

    except HttpError as error:
        error_details = error.content.decode('utf-8')
        print(f"Google API Error in export_grades_to_sheet: {error.resp.status} - {error_details}")
        return JSONResponse(
            content={"error": f"Google API Error: {error_details}"},
            status_code=error.resp.status
        )
    except Exception as e:
        print(f"Unexpected error in export_grades_to_sheet: {e}")
        return JSONResponse(
            content={"error": f"An unexpected error occurred: {str(e)}"},
            status_code=500
        )


# --- 9. ANALYTICS ENDPOINTS ---

@app.get('/api/analytics/distribution')
async def get_grade_distribution(
    request: Request,
    course_id: str = None,
    assignment_id: str = None,
    student_name: str = None
):
    """
    Get grade distribution for histogram/bar chart.
    Can filter by course, assignment, or student.
    """
    try:
        # Build query filter
        query = {}
        if course_id:
            query['course_id'] = course_id
        if assignment_id:
            query['assignment_id'] = assignment_id
        if student_name:
            query['student_name'] = student_name
        
        # Get all grades matching the filter
        if grades_collection is not None:
            grades = list(grades_collection.find(query, {'assignedGrade': 1, '_id': 0}))
        else:
            # Fallback to in-memory
            grades = [
                {'assignedGrade': g['assignedGrade']}
                for g in graded_assignments_history
                if (not course_id or g.get('course_id') == course_id) and
                   (not assignment_id or g.get('assignment_id') == assignment_id) and
                   (not student_name or g.get('student_name') == student_name)
            ]
        
        if not grades:
            return {
                "distribution": {"0-50": 0, "51-70": 0, "71-85": 0, "86-100": 0},
                "total_graded": 0,
                "average_grade": 0
            }
        
        # Calculate distribution
        distribution = {"0-50": 0, "51-70": 0, "71-85": 0, "86-100": 0}
        total_score = 0
        
        for grade_doc in grades:
            grade = grade_doc['assignedGrade']
            total_score += grade
            
            if grade <= 50:
                distribution["0-50"] += 1
            elif grade <= 70:
                distribution["51-70"] += 1
            elif grade <= 85:
                distribution["71-85"] += 1
            else:
                distribution["86-100"] += 1
        
        return {
            "distribution": distribution,
            "total_graded": len(grades),
            "average_grade": round(total_score / len(grades), 2)
        }
    except Exception as e:
        print(f"❌ Error in get_grade_distribution: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


@app.get('/api/analytics/student-history/{student_name}')
async def get_student_history(student_name: str, course_id: str = None):
    """
    Get all previous grades and performance history for a specific student.
    Can optionally filter by course.
    """
    try:
        query = {'student_name': student_name}
        if course_id:
            query['course_id'] = course_id
        
        # Get all grades for this student
        if grades_collection is not None:
            grades = list(grades_collection.find(
                query,
                {'_id': 0}
            ).sort('timestamp', -1))
        else:
            # Fallback to in-memory
            grades = [
                g for g in graded_assignments_history
                if g.get('student_name') == student_name and
                   (not course_id or g.get('course_id') == course_id)
            ]
            grades.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        
        if not grades:
            return {
                "student_name": student_name,
                "grades": [],
                "average_grade": 0,
                "total_assignments": 0,
                "courses": [],
                "performance_trend": []
            }
        
        # Calculate statistics
        total_score = sum(g['assignedGrade'] for g in grades)
        average_grade = round(total_score / len(grades), 2)
        
        # Get unique courses
        courses = list(set(g['course_name'] for g in grades))
        
        # Performance trend (last 10 assignments)
        performance_trend = [
            {
                "assignment": g['assignment_title'],
                "grade": g['assignedGrade'],
                "date": g['timestamp']
            }
            for g in grades[:10]
        ]
        
        # Grade by course
        course_performance = {}
        for grade in grades:
            course = grade['course_name']
            if course not in course_performance:
                course_performance[course] = []
            course_performance[course].append(grade['assignedGrade'])
        
        course_averages = {
            course: round(sum(grades_list) / len(grades_list), 2)
            for course, grades_list in course_performance.items()
        }
        
        return {
            "student_name": student_name,
            "grades": grades,
            "average_grade": average_grade,
            "total_assignments": len(grades),
            "courses": courses,
            "performance_trend": performance_trend,
            "course_averages": course_averages
        }
    except Exception as e:
        print(f"❌ Error in get_student_history: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


@app.get('/api/analytics/course-stats/{course_id}')
async def get_course_stats(course_id: str):
    """
    Get comprehensive statistics for a specific course.
    Includes assignment-wise breakdown and student performance.
    """
    try:
        # Use MongoDB aggregation if available
        if grades_collection is not None:
            pipeline = [
                {"$match": {"course_id": course_id}},
                {"$group": {
                    "_id": "$assignment_id",
                    "assignment_title": {"$first": "$assignment_title"},
                    "course_name": {"$first": "$course_name"},
                    "avg_grade": {"$avg": "$assignedGrade"},
                    "min_grade": {"$min": "$assignedGrade"},
                    "max_grade": {"$max": "$assignedGrade"},
                    "count": {"$sum": 1},
                    "grades": {"$push": "$assignedGrade"}
                }},
                {"$sort": {"avg_grade": -1}}
            ]
            
            stats = list(grades_collection.aggregate(pipeline))
        else:
            # Fallback: manual aggregation for in-memory
            course_grades = [g for g in graded_assignments_history if g.get('course_id') == course_id]
            if not course_grades:
                return {
                    "course_id": course_id,
                    "course_name": "Unknown",
                    "total_assignments": 0,
                    "total_graded": 0,
                    "overall_average": 0,
                    "assignments": []
                }
            
            # Group by assignment
            assignment_map = {}
            for g in course_grades:
                aid = g['assignment_id']
                if aid not in assignment_map:
                    assignment_map[aid] = {
                        '_id': aid,
                        'assignment_title': g['assignment_title'],
                        'course_name': g['course_name'],
                        'grades': []
                    }
                assignment_map[aid]['grades'].append(g['assignedGrade'])
            
            stats = []
            for aid, data in assignment_map.items():
                grades_list = data['grades']
                stats.append({
                    '_id': aid,
                    'assignment_title': data['assignment_title'],
                    'course_name': data['course_name'],
                    'avg_grade': sum(grades_list) / len(grades_list),
                    'min_grade': min(grades_list),
                    'max_grade': max(grades_list),
                    'count': len(grades_list),
                    'grades': grades_list
                })
            stats.sort(key=lambda x: x['avg_grade'], reverse=True)
        
        if not stats:
            return {
                "course_id": course_id,
                "course_name": "Unknown",
                "total_assignments": 0,
                "total_graded": 0,
                "overall_average": 0,
                "assignments": []
            }
        
        # Calculate overall statistics
        total_graded = sum(s['count'] for s in stats)
        all_grades = []
        for s in stats:
            all_grades.extend(s['grades'])
        
        overall_average = round(sum(all_grades) / len(all_grades), 2) if all_grades else 0
        
        # Format assignment stats
        formatted_stats = []
        for stat in stats:
            # Calculate standard deviation
            grades = stat['grades']
            mean = stat['avg_grade']
            variance = sum((g - mean) ** 2 for g in grades) / len(grades)
            std_dev = round(variance ** 0.5, 2)
            
            formatted_stats.append({
                "assignment_id": stat["_id"],
                "assignment_title": stat["assignment_title"],
                "average_grade": round(stat["avg_grade"], 2),
                "min_grade": stat["min_grade"],
                "max_grade": stat["max_grade"],
                "submissions_count": stat["count"],
                "std_deviation": std_dev
            })
        
        return {
            "course_id": course_id,
            "course_name": stats[0]['course_name'],
            "total_assignments": len(stats),
            "total_graded": total_graded,
            "overall_average": overall_average,
            "assignments": formatted_stats
        }
    except Exception as e:
        print(f"❌ Error in get_course_stats: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


@app.get('/api/analytics/students')
async def get_all_students(course_id: str = None, sort_by: str = "average_grade"):
    """
    Get list of all students with their performance metrics.
    Can filter by course and sort by various fields.
    """
    try:
        # Build match query
        match_query = {}
        if course_id:
            match_query['course_id'] = course_id
        
        # Use MongoDB aggregation if available
        if grades_collection is not None:
            pipeline = [
                {"$match": match_query},
                {"$group": {
                    "_id": "$student_name",
                    "avg_grade": {"$avg": "$assignedGrade"},
                    "total_assignments": {"$sum": 1},
                    "highest_grade": {"$max": "$assignedGrade"},
                    "lowest_grade": {"$min": "$assignedGrade"},
                    "courses": {"$addToSet": "$course_name"},
                    "recent_grade": {"$last": "$assignedGrade"},
                    "recent_assignment": {"$last": "$assignment_title"}
                }},
                {"$sort": {"avg_grade": -1 if sort_by == "average_grade" else 1}}
            ]
            
            students = list(grades_collection.aggregate(pipeline))
        else:
            # Fallback: manual aggregation
            filtered_grades = [
                g for g in graded_assignments_history
                if not course_id or g.get('course_id') == course_id
            ]
            
            student_map = {}
            for g in filtered_grades:
                sname = g['student_name']
                if sname not in student_map:
                    student_map[sname] = {
                        '_id': sname,
                        'grades': [],
                        'courses': set(),
                        'recent_grade': g['assignedGrade'],
                        'recent_assignment': g['assignment_title']
                    }
                student_map[sname]['grades'].append(g['assignedGrade'])
                student_map[sname]['courses'].add(g['course_name'])
            
            students = []
            for sname, data in student_map.items():
                grades = data['grades']
                students.append({
                    '_id': sname,
                    'avg_grade': sum(grades) / len(grades),
                    'total_assignments': len(grades),
                    'highest_grade': max(grades),
                    'lowest_grade': min(grades),
                    'courses': list(data['courses']),
                    'recent_grade': data['recent_grade'],
                    'recent_assignment': data['recent_assignment']
                })
            students.sort(key=lambda x: x['avg_grade'], reverse=True)
        
        # Format for frontend
        formatted_students = []
        for student in students:
            formatted_students.append({
                "student_name": student["_id"],
                "average_grade": round(student["avg_grade"], 2),
                "total_assignments": student["total_assignments"],
                "highest_grade": student["highest_grade"],
                "lowest_grade": student["lowest_grade"],
                "courses": student["courses"],
                "recent_performance": {
                    "grade": student["recent_grade"],
                    "assignment": student["recent_assignment"]
                }
            })
        
        return {
            "total_students": len(formatted_students),
            "students": formatted_students
        }
    except Exception as e:
        print(f"❌ Error in get_all_students: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


@app.get('/api/analytics/compare')
async def compare_performance(
    type: str,  # "courses" or "assignments"
    ids: str = None  # comma-separated IDs to compare
):
    """
    Compare performance across multiple courses or assignments.
    Returns comparative statistics for visualization.
    """
    try:
        if type == "courses":
            # Compare multiple courses
            course_ids = ids.split(',') if ids else []
            
            comparison_data = []
            for course_id in course_ids:
                if grades_collection is not None:
                    grades = list(grades_collection.find(
                        {"course_id": course_id},
                        {"assignedGrade": 1, "course_name": 1, "_id": 0}
                    ))
                else:
                    grades = [
                        {"assignedGrade": g['assignedGrade'], "course_name": g['course_name']}
                        for g in graded_assignments_history
                        if g.get('course_id') == course_id
                    ]
                
                if grades:
                    total_score = sum(g['assignedGrade'] for g in grades)
                    comparison_data.append({
                        "course_id": course_id,
                        "course_name": grades[0]['course_name'],
                        "average_grade": round(total_score / len(grades), 2),
                        "total_graded": len(grades)
                    })
            
            return {
                "type": "courses",
                "data": comparison_data
            }
        
        elif type == "assignments":
            # Compare multiple assignments
            assignment_ids = ids.split(',') if ids else []
            
            comparison_data = []
            for assignment_id in assignment_ids:
                if grades_collection is not None:
                    grades = list(grades_collection.find(
                        {"assignment_id": assignment_id},
                        {"assignedGrade": 1, "assignment_title": 1, "_id": 0}
                    ))
                else:
                    grades = [
                        {"assignedGrade": g['assignedGrade'], "assignment_title": g['assignment_title']}
                        for g in graded_assignments_history
                        if g.get('assignment_id') == assignment_id
                    ]
                
                if grades:
                    total_score = sum(g['assignedGrade'] for g in grades)
                    comparison_data.append({
                        "assignment_id": assignment_id,
                        "assignment_title": grades[0]['assignment_title'],
                        "average_grade": round(total_score / len(grades), 2),
                        "total_graded": len(grades)
                    })
            
            return {
                "type": "assignments",
                "data": comparison_data
            }
        
        else:
            return JSONResponse(
                content={"error": "Invalid comparison type. Use 'courses' or 'assignments'."},
                status_code=400
            )
    
    except Exception as e:
        print(f"❌ Error in compare_performance: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


@app.get('/api/analytics/trends')
async def get_performance_trends(
    course_id: str = None,
    student_name: str = None,
    time_period: str = "all"  # "week", "month", "semester", "all"
):
    """
    Get performance trends over time.
    Shows how grades have changed over different time periods.
    """
    try:
        from datetime import timedelta
        
        # Build query
        query = {}
        if course_id:
            query['course_id'] = course_id
        if student_name:
            query['student_name'] = student_name
        
        # Add time filter
        if time_period != "all":
            now = datetime.datetime.now()
            if time_period == "week":
                start_date = now - timedelta(days=7)
            elif time_period == "month":
                start_date = now - timedelta(days=30)
            elif time_period == "semester":
                start_date = now - timedelta(days=120)
            else:
                start_date = datetime.datetime(2000, 1, 1)
            
            query['timestamp'] = {"$gte": start_date.isoformat()}
        
        # Get grades sorted by time
        if grades_collection is not None:
            grades = list(grades_collection.find(
                query,
                {'assignedGrade': 1, 'timestamp': 1, 'assignment_title': 1, '_id': 0}
            ).sort('timestamp', 1))
        else:
            # Fallback to in-memory
            grades = [
                {'assignedGrade': g['assignedGrade'], 'timestamp': g['timestamp'], 
                 'assignment_title': g['assignment_title']}
                for g in graded_assignments_history
                if (not course_id or g.get('course_id') == course_id) and
                   (not student_name or g.get('student_name') == student_name)
            ]
            grades.sort(key=lambda x: x['timestamp'])
        
        if not grades:
            return {
                "trend_data": [],
                "overall_trend": "no_data"
            }
        
        # Format trend data
        trend_data = [
            {
                "date": g['timestamp'],
                "grade": g['assignedGrade'],
                "assignment": g.get('assignment_title', 'Unknown')
            }
            for g in grades
        ]
        
        # Calculate trend direction
        if len(grades) >= 3:
            first_third_avg = sum(g['assignedGrade'] for g in grades[:len(grades)//3]) / (len(grades)//3)
            last_third_avg = sum(g['assignedGrade'] for g in grades[-len(grades)//3:]) / (len(grades)//3)
            
            if last_third_avg > first_third_avg + 5:
                overall_trend = "improving"
            elif last_third_avg < first_third_avg - 5:
                overall_trend = "declining"
            else:
                overall_trend = "stable"
        else:
            overall_trend = "insufficient_data"
        
        return {
            "trend_data": trend_data,
            "overall_trend": overall_trend,
            "total_data_points": len(grades)
        }
    
    except Exception as e:
        print(f"❌ Error in get_performance_trends: {e}")
        return JSONResponse(
            content={"error": str(e)},
            status_code=500
        )


# --- DATABASE STATUS ENDPOINT ---

@app.get('/api/db_status')
async def get_db_status():
    """
    Returns the current database connection status.
    """
    return {
        "mongodb_connected": grades_collection is not None,
        "storage_type": "MongoDB" if grades_collection is not None else "In-Memory",
        "total_grades": len(graded_assignments_history) if grades_collection is None else "Check MongoDB",
        "database_name": db.name if db is not None else None
    }


# --- 6. MAIN APPLICATION RUNNER ---

# CHANGED: Replaced Flask's 'app.run' with 'uvicorn.run'
if __name__ == '__main__':
    # Flask's 'debug=True' is similar to 'reload=True' in uvicorn,
    # but 'reload=True' must be run from the command line.
    # This is the direct equivalent of 'app.run(port=8000)'
    # To run with reload (like Flask debug): uvicorn main:app --port 8000 --reload
    # (Assuming your file is named 'main.py')
    uvicorn.run(app, host="127.0.0.1", port=8000)